#!/usr/bin/env python3
"""
Generate the OrgPulse product overview Word document (client + investor facing).

Run with the project venv that has python-docx installed:
    ./.docbuild-venv/bin/python scripts/generate_product_doc.py

Output: OrgPulse-Product-Overview.docx at the repo root (~5 pages).
The document is product/tech focused (no market sizing or pricing) with
qualitative ROI, brand styling, diagrams/flows, persona cards, and
[Screenshot placeholder] blocks the user fills in later.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette ──────────────────────────────────────────────────────────
BLUE = "0176D3"        # Salesforce blue (primary accent)
NAVY = "0B1B33"        # deep navy for cover
DARK = "1A1A2E"        # heading text
GREY = "5A6473"        # muted body
LIGHT = "F2F5FA"       # light panel
BORDER = "C9D4E2"      # box borders
GREEN = "2E844A"
AMBER = "FE9339"
RED = "EA001E"
WHITE = "FFFFFF"

VERSION = "v1.11.0"
DATE = "June 2026"


# ── Low-level OOXML helpers ────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def set_cell_borders(cell, color=BORDER, sz=6, sides=("top", "bottom", "start", "end")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in sides:
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def set_row_height(row, pts):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(pts * 20)))
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)


def shade_paragraph(par, hex_color):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


# ── Text helpers ───────────────────────────────────────────────────────────
def run(par, text, size=10, bold=False, color=DARK, italic=False, font="Segoe UI"):
    r = par.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.name = font
    return r


def cell_text(cell, text, size=9.5, bold=False, color=DARK, align="left",
              valign="center", italic=False):
    cell.vertical_alignment = {
        "center": WD_ALIGN_VERTICAL.CENTER,
        "top": WD_ALIGN_VERTICAL.TOP,
    }[valign]
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run(p, text, size=size, bold=bold, color=color, italic=italic)
    return p


def heading(doc, text, num=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    if num:
        run(p, f"{num}  ", size=15, bold=True, color=BLUE)
    run(p, text, size=15, bold=True, color=DARK)
    # accent underline
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), BLUE)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def body(doc, text, size=10, color=DARK, space_after=4, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run(p, text, size=size, color=color, bold=bold)
    return p


def bullet(doc, label, text, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(10)
    run(p, "▸  ", size=10, bold=True, color=BLUE)
    if label:
        run(p, f"{label} ", size=10, bold=True, color=color)
    run(p, text, size=10, color=GREY)
    return p


def spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)
    return p


def screenshot_placeholder(doc, caption, height_pts=120):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, LIGHT)
    set_cell_borders(cell, color=BLUE, sz=6)
    set_row_height(t.rows[0], height_pts)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "🖼  ", size=12, color=BLUE)
    run(p, f"[ Screenshot placeholder — {caption} ]", size=10, italic=True, color=GREY)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, "Paste image here", size=8, italic=True, color=BORDER)
    spacer(doc, 4)
    return t


# ── Document setup ─────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(DARK)

    # Tight margins to fit 5 pages
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_footer(doc)

    # ===================== PAGE 1 — COVER + EXEC SUMMARY =====================
    cover(doc)
    exec_summary(doc)

    # ===================== PAGE 2 — PROBLEM & WHO IT'S FOR ===================
    doc.add_page_break()
    problem_and_audience(doc)

    # ===================== PAGE 3 — FEATURES ================================
    doc.add_page_break()
    features(doc)

    # ===================== PAGE 4 — ARCHITECTURE ============================
    doc.add_page_break()
    architecture(doc)

    # ===================== PAGE 5 — WHY / VALUE / GET STARTED ===============
    doc.add_page_break()
    why_value_start(doc)

    out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                       "OrgPulse-Product-Overview.docx")
    doc.save(out)
    return out


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "OrgPulse — Confidential Product Overview    ·    ", size=8, color=GREY)
    # page number field
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16"); rpr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), GREY); rpr.append(color)
    r.append(rpr)
    t = OxmlElement("w:t"); t.text = "1"; r.append(t)
    fldSimple.append(r)
    p._p.append(fldSimple)


# ── PAGE 1 ─────────────────────────────────────────────────────────────────
def cover(doc):
    band = doc.add_table(rows=1, cols=1)
    no_table_borders(band)
    set_row_height(band.rows[0], 96)
    c = band.rows[0].cells[0]
    set_cell_bg(c, NAVY)
    set_cell_margins(c, top=200, bottom=200, left=260, right=260)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]
    run(p, "OrgPulse", size=30, bold=True, color=WHITE)
    p2 = c.add_paragraph()
    run(p2, "Salesforce Architecture Health & Insights — right inside VS Code", size=12, color="BFD7F2")
    p3 = c.add_paragraph()
    p3.paragraph_format.space_before = Pt(6)
    run(p3, f"Product Overview   ·   {VERSION}   ·   {DATE}", size=9.5, color="8FB6E6")
    spacer(doc, 8)


def exec_summary(doc):
    heading(doc, "Executive Summary")
    body(doc,
         "OrgPulse turns a Salesforce org's hidden complexity into a clear, continuous health signal. "
         "It connects to an authenticated org through the Salesforce CLI, analyses metadata across 17 "
         "architecture dimensions, scores overall health 0–100 with an A–F grade and run-over-run trend, "
         "and surfaces prioritised, actionable findings — all without leaving the developer's IDE.",
         space_after=4)
    body(doc,
         "Beyond static analysis, OrgPulse adds an AI layer that is model-agnostic: it works with any model "
         "available through VS Code (Claude, GPT, Codex, Gemini) or a fully local model (Ollama / LM Studio) "
         "with no cloud dependency. Architects can generate a boardroom-grade CTA Architecture Review, or use "
         "“Ask the Architect” to query the live org in plain language. Everything runs metadata-only and "
         "local-first, so sensitive data never leaves the environment without explicit consent.",
         space_after=6)
    screenshot_placeholder(doc, "OrgPulse health dashboard (overview + score ring)", height_pts=150)


# ── PAGE 2 ─────────────────────────────────────────────────────────────────
def problem_and_audience(doc):
    heading(doc, "The Problem", num="01")
    body(doc,
         "Salesforce orgs accrete complexity faster than any team can manually track. Years of clicks and "
         "code leave behind systemic risk that traditional reviews miss:", space_after=4)
    for lab, txt in [
        ("Hidden technical debt —", "trigger sprawl, SOQL/DML in loops, and untested code accumulate silently."),
        ("Security & performance risk —", "over-broad permissions, exposed credentials, and non-selective queries surface in production."),
        ("No continuous visibility —", "architecture health is a point-in-time PDF, not a living signal teams can act on."),
        ("Expensive expert reviews —", "a Certified Technical Architect engagement is costly and infrequent — gaps go unseen between reviews."),
    ]:
        bullet(doc, lab, txt)
    spacer(doc, 2)

    # Before / After table
    heading(doc, "The Solution — Before → After")
    rows = [
        ("Manual reviews miss systemic patterns", "Automated analysis of every class, trigger, Flow & object"),
        ("Security gaps lurk in permissions", "Dedicated security scanner flags risks instantly"),
        ("Tech debt builds silently", "Quantified health score with run-over-run trend"),
        ("SOQL issues found in production", "Live query selectivity (EXPLAIN) before deploy"),
        ("CTA reviews are slow and periodic", "CTA-grade AI review on demand, in minutes"),
    ]
    t = doc.add_table(rows=len(rows) + 1, cols=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = t.rows[0].cells
    for i, (head, col) in enumerate([("Without OrgPulse", RED), ("With OrgPulse", GREEN)]):
        set_cell_bg(hdr[i], col)
        set_cell_borders(hdr[i], color=WHITE, sz=4)
        cell_text(hdr[i], head, bold=True, color=WHITE, align="center", size=10)
    for ri, (left, right) in enumerate(rows, start=1):
        for ci, (txt, fill) in enumerate([(left, "FCE9EA"), (right, "E7F3EB")]):
            cell = t.rows[ri].cells[ci]
            set_cell_bg(cell, fill)
            set_cell_borders(cell, color=WHITE, sz=4)
            set_cell_margins(cell)
            cell_text(cell, txt, size=9, color=DARK)
    spacer(doc, 6)

    # Beneficiaries + customers
    heading(doc, "Who Benefits & Who Buys")
    personas = [
        ("👩‍💻", "Developers", "Inline issues with file/line, AI fixes, quick wins."),
        ("🏛️", "Architects", "CTA review, dependency graph, LDV & governor risk."),
        ("📈", "Tech Leads / Mgrs", "Grade, trend, and risk at a glance — no code needed."),
        ("🎓", "CTAs / Reviewers", "Repeatable, evidence-based architecture assessments."),
    ]
    pt = doc.add_table(rows=1, cols=4)
    no_table_borders(pt)
    for i, (icon, title, desc) in enumerate(personas):
        c = pt.rows[0].cells[i]
        set_cell_bg(c, LIGHT)
        set_cell_borders(c, color=BORDER, sz=4)
        set_cell_margins(c, top=80, bottom=80, left=110, right=110)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, icon, size=16)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p2, title, size=9.5, bold=True, color=BLUE)
        p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        run(p3, desc, size=8, color=GREY)
    spacer(doc, 4)
    body(doc, "End customers:", bold=True, space_after=2)
    bullet(doc, "Salesforce SIs & consulting partners —", "standardise delivery quality and health hand-offs across clients.")
    bullet(doc, "AppExchange ISVs —", "pass Security Review faster and keep managed packages clean.")
    bullet(doc, "Large-org enterprises —", "continuous architecture governance for high-volume, mission-critical orgs.")
    bullet(doc, "Managed service providers —", "monitor many orgs with repeatable, exportable scorecards.")


# ── PAGE 3 ─────────────────────────────────────────────────────────────────
def features(doc):
    heading(doc, "Features & Capabilities", num="02")

    grid = [
        ("🔬 17-Analyzer Engine",
         "Apex quality, automation design, data model, query performance, test coverage, security & "
         "permissions, integrations, governor limits, LWC, dependencies, technical debt, CTA architecture & more."),
        ("🧠 Model-Agnostic AI",
         "Use any VS Code language model (Claude, GPT, Codex, Gemini) or a fully local model "
         "(Ollama / LM Studio) — no Copilot lock-in, no cloud requirement."),
        ("💬 Ask the Architect",
         "Tool-augmented, read-only live Q&A — the AI runs SOQL / limits / EXPLAIN against your org to "
         "answer architecture questions grounded in real data."),
        ("🎯 CTA Architecture Review",
         "On-demand, boardroom-grade AI report: verdict, maturity, risk heatmap, domain findings, and "
         "prioritised recommendations."),
        ("📊 Interactive Dashboard",
         "Health score ring + A–F grade, run-over-run trend, drill-down per issue, live governor limits, "
         "and domain tabs."),
        ("📤 Multi-Format Export",
         "HTML (stakeholders), JSON (tooling), SARIF 2.1.0 (GitHub code scanning / PR checks), CSV, and PDF."),
    ]
    t = doc.add_table(rows=3, cols=2)
    no_table_borders(t)
    for idx, (title, desc) in enumerate(grid):
        cell = t.rows[idx // 2].cells[idx % 2]
        set_cell_bg(cell, LIGHT)
        set_cell_borders(cell, color=BORDER, sz=4)
        set_cell_margins(cell, top=90, bottom=90, left=130, right=130)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        run(p, title, size=10.5, bold=True, color=BLUE)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        run(p2, desc, size=8.5, color=GREY)
    spacer(doc, 6)

    heading(doc, "Security-First by Design")
    modes = [
        ("🟢 Safe", GREEN, "100% local — no AI, no network traffic."),
        ("🟡 Standard", AMBER, "AI on aggregated insights only — no raw metadata sent."),
        ("🔴 Advanced", RED, "Deeper AI reasoning with explicit per-session consent."),
    ]
    mt = doc.add_table(rows=1, cols=3)
    no_table_borders(mt)
    for i, (title, col, desc) in enumerate(modes):
        c = mt.rows[0].cells[i]
        set_cell_bg(c, LIGHT)
        set_cell_borders(c, color=col, sz=8)
        set_cell_margins(c, top=70, bottom=70, left=110, right=110)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, title, size=10, bold=True, color=col)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        run(p2, desc, size=8, color=GREY)
    spacer(doc, 4)
    body(doc,
         "Metadata-only · no record data or PII accessed · nothing stored outside your environment · "
         "AI is consent-gated and, with a local model, fully offline.", size=9, color=GREY, space_after=6)

    screenshot_placeholder(doc, "CTA Architecture Review report", height_pts=92)
    screenshot_placeholder(doc, "Ask the Architect — live org Q&A", height_pts=92)


# ── PAGE 4 ─────────────────────────────────────────────────────────────────
def diagram_box(cell, title, sub, fill, border, title_color=WHITE):
    set_cell_bg(cell, fill)
    set_cell_borders(cell, color=border, sz=8)
    set_cell_margins(cell, top=70, bottom=70, left=120, right=120)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, title, size=10, bold=True, color=title_color)
    if sub:
        p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        run(p2, sub, size=8, color=title_color)


def connector(doc, text="▼"):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    run(p, text, size=11, bold=True, color=BLUE)


def architecture(doc):
    heading(doc, "Architecture & How It Works", num="03")

    # Layered architecture diagram (3 stacked layers)
    layers = [
        ("VS Code Extension Host  (TypeScript, esbuild bundle)",
         "Analyzers → Rules Engine → Health Score → Webview Dashboard", BLUE, BLUE),
        ("AI Layer  ·  model-agnostic",
         "VS Code Language Model API (Claude / GPT / Gemini)  ·  Local OpenAI-compatible endpoint (Ollama / LM Studio)  ·  read-only org tools",
         "3B5BDB", "3B5BDB"),
        ("Salesforce Org  ·  via Salesforce CLI (sf)",
         "Tooling & Data API — metadata, EXPLAIN, limits  (read-only)", NAVY, NAVY),
    ]
    for title, sub, fill, border in layers:
        t = doc.add_table(rows=1, cols=1)
        no_table_borders(t)
        diagram_box(t.rows[0].cells[0], title, sub, fill, border)
        if (title, sub, fill, border) != layers[-1]:
            connector(doc, "▲  ▼")
    spacer(doc, 6)

    # Analysis flow (numbered horizontal-ish)
    heading(doc, "Analysis Flow")
    steps = [
        ("1 Connect", "Authenticate via sf CLI"),
        ("2 Fetch", "Pull org metadata"),
        ("3 Analyse", "17 analyzers + rules"),
        ("4 Score", "0–100, A–F, trend"),
        ("5 Review", "Dashboard · CTA · Q&A"),
        ("6 Export", "HTML · SARIF · PDF"),
    ]
    ft = doc.add_table(rows=1, cols=len(steps) * 2 - 1)
    no_table_borders(ft)
    col = 0
    for i, (title, sub) in enumerate(steps):
        c = ft.rows[0].cells[col]
        set_cell_bg(c, LIGHT)
        set_cell_borders(c, color=BLUE, sz=4)
        set_cell_margins(c, top=60, bottom=60, left=50, right=50)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, title, size=8.5, bold=True, color=BLUE)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        run(p2, sub, size=7, color=GREY)
        col += 1
        if i < len(steps) - 1:
            arrow = ft.rows[0].cells[col]
            arrow.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            ap = arrow.paragraphs[0]; ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(ap, "→", size=11, bold=True, color=BLUE)
            col += 1
    spacer(doc, 6)

    # Tech stack + security boundary side by side
    heading(doc, "Tech Stack & Trust Boundary")
    two = doc.add_table(rows=1, cols=2)
    no_table_borders(two)
    # left: tech stack
    left = two.rows[0].cells[0]
    set_cell_margins(left, right=160)
    lp = left.paragraphs[0]
    run(lp, "Technology", size=10, bold=True, color=DARK)
    for lab, txt in [
        ("Language —", "TypeScript (strict), ES2022"),
        ("Build —", "esbuild single bundle; zero runtime dependencies"),
        ("Platform —", "VS Code Extension API + Webview dashboard"),
        ("Org access —", "Salesforce CLI (sf) — Tooling & Data API"),
        ("AI —", "VS Code LM API + local OpenAI-compatible endpoint"),
        ("CI-ready —", "SARIF 2.1.0 export for code scanning"),
    ]:
        bp = left.add_paragraph(); bp.paragraph_format.space_after = Pt(1)
        run(bp, "• ", size=9, bold=True, color=BLUE)
        run(bp, lab + " ", size=9, bold=True, color=DARK)
        run(bp, txt, size=9, color=GREY)
    # right: trust boundary box
    right = two.rows[0].cells[1]
    set_cell_bg(right, "E7F3EB")
    set_cell_borders(right, color=GREEN, sz=8)
    set_cell_margins(right, top=110, bottom=110, left=150, right=150)
    rp = right.paragraphs[0]
    run(rp, "🔒 Trust Boundary", size=10, bold=True, color=GREEN)
    for txt in [
        "Metadata only — no records, no PII",
        "Local-first engine — offline capable",
        "Nothing stored outside your machine",
        "AI consent-gated; local model = 0 egress",
    ]:
        q = right.add_paragraph(); q.paragraph_format.space_after = Pt(1)
        run(q, "✔ ", size=9, bold=True, color=GREEN)
        run(q, txt, size=8.5, color=DARK)


# ── PAGE 5 ─────────────────────────────────────────────────────────────────
def why_value_start(doc):
    heading(doc, "Why OrgPulse — and Why Only OrgPulse", num="04")
    diffs = [
        ("In the IDE, not a portal", "Analysis, AI, and fixes where developers already work — zero context switching."),
        ("Local-first & metadata-only", "No record data or PII; runs offline. Enterprise- and Security-Review-friendly by default."),
        ("Truly model-agnostic", "Any VS Code model or a fully local LLM — the only Salesforce health tool with no AI vendor lock-in."),
        ("CTA-grade, on demand", "Repeatable architecture reviews in minutes — not a periodic, costly human engagement."),
        ("Live, tool-augmented Q&A", "“Ask the Architect” queries the real org (read-only) — answers grounded in data, not guesses."),
        ("CI/CD native", "SARIF export plugs straight into GitHub code scanning and PR checks."),
    ]
    t = doc.add_table(rows=3, cols=2)
    no_table_borders(t)
    for idx, (title, desc) in enumerate(diffs):
        cell = t.rows[idx // 2].cells[idx % 2]
        set_cell_bg(cell, LIGHT)
        set_cell_borders(cell, color=BLUE, sz=4)
        set_cell_margins(cell, top=70, bottom=70, left=120, right=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        run(p, "✦ " + title, size=9.5, bold=True, color=BLUE)
        p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(0)
        run(p2, desc, size=8.5, color=GREY)
    spacer(doc, 6)

    heading(doc, "Business Value & ROI")
    for lab, txt in [
        ("Continuous, not point-in-time —", "every run is a fresh scorecard with trend, so risk is caught between formal reviews."),
        ("Minutes vs weeks —", "a CTA-grade assessment on demand instead of a multi-week expert engagement (replace with [your benchmark])."),
        ("Shift-left risk —", "security, performance, and governor issues found pre-deploy — [your % defects caught earlier]."),
        ("Standardised delivery —", "partners & MSPs apply one consistent health bar across every org — [your orgs/month]."),
    ]:
        bullet(doc, lab, txt)
    spacer(doc, 4)

    heading(doc, "Get Started in Minutes")
    qs = [
        ("Install", "Add the OrgPulse VSIX (or Marketplace) to VS Code 1.100+."),
        ("Authenticate", "Connect an org: sf org login web."),
        ("Analyse", "Run “OrgPulse: Run Org Health Analysis”."),
        ("Act & Share", "Review the dashboard, ask the architect, export HTML/SARIF/PDF."),
    ]
    qt = doc.add_table(rows=1, cols=len(qs))
    no_table_borders(qt)
    for i, (title, desc) in enumerate(qs):
        c = qt.rows[0].cells[i]
        set_cell_bg(c, NAVY)
        set_cell_borders(c, color=WHITE, sz=4)
        set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, f"{i+1}. {title}", size=9.5, bold=True, color=WHITE)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        run(p2, desc, size=7.5, color="BFD7F2")
    spacer(doc, 6)

    heading(doc, "Roadmap (Indicative)")
    bullet(doc, "Near-term —", "scheduled scans & CI gate action; richer trend history & diffing.")
    bullet(doc, "Mid-term —", "team dashboards / multi-org portfolio view; custom rule marketplace.")
    bullet(doc, "Exploration —", "auto-fix PRs, deeper LWC analysis, and partner reporting templates.")
    spacer(doc, 6)

    # closing band
    band = doc.add_table(rows=1, cols=1)
    no_table_borders(band)
    c = band.rows[0].cells[0]
    set_cell_bg(c, BLUE)
    set_cell_margins(c, top=120, bottom=120, left=160, right=160)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "Bring continuous architecture intelligence to every Salesforce org.", size=11, bold=True, color=WHITE)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    run(p2, "OrgPulse  ·  [contact email]  ·  [website]  ·  " + VERSION, size=9, color="DCEBFB")


if __name__ == "__main__":
    path = build()
    print("Wrote:", path)
